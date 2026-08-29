"""
Unified Document Loader for PDF, TXT, and DOCX Indian Legal Documents.
"""

from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import logging
from src.document.models import Document, PageContent, TextBlock
from src.document.pdf_extractor import PDFExtractor

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


class DocumentLoader:
    """
    Loads legal case documents from disk across multiple formats.
    Ensures consistent Document schema preserving page boundaries and block layouts.
    """

    def __init__(
        self,
        ocr_enabled: bool = True,
        scanned_threshold_chars: int = 50,
        ocr_languages: Optional[List[str]] = None,
    ):
        self.pdf_extractor = PDFExtractor(
            scanned_threshold_chars=scanned_threshold_chars,
            ocr_enabled=ocr_enabled,
            ocr_languages=ocr_languages,
        )

    def load(self, file_path: Union[str, Path], document_id: Optional[str] = None) -> Document:
        """
        Load a document and return a structured Document object.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Case document not found at: {path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            supported_str = ", ".join(SUPPORTED_EXTENSIONS)
            raise ValueError(
                f"Unsupported document format '{ext}' for file '{path.name}'. "
                f"Supported formats are: {supported_str}"
            )

        doc_id = document_id or path.stem

        if ext == ".pdf":
            return self.pdf_extractor.extract(path, document_id=doc_id)
        elif ext == ".txt":
            return self._load_txt(path, document_id=doc_id)
        elif ext == ".docx":
            return self._load_docx(path, document_id=doc_id)
        else:
            raise ValueError(f"Unhandled extension: {ext}")

    def _load_txt(self, path: Path, document_id: str) -> Document:
        """Load text document, splitting on form feeds (\\f) or page markers if present."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        # Check for form feed or standard page markers
        if "\x0c" in content:
            raw_pages = content.split("\x0c")
        elif "--- PAGE BREAK ---" in content:
            raw_pages = content.split("--- PAGE BREAK ---")
        elif "\n\n[PAGE " in content:
            import re
            raw_pages = re.split(r"\n\n\[PAGE \d+\]\n", content)
        else:
            # Single page or synthetic document
            raw_pages = [content]

        pages: List[PageContent] = []
        for idx, page_raw in enumerate(raw_pages):
            page_text = page_raw.strip()
            if not page_text and len(raw_pages) > 1:
                continue

            blocks = []
            paragraphs = [p.strip() for p in page_text.split("\n\n") if p.strip()]
            for b_idx, para in enumerate(paragraphs):
                blocks.append(
                    TextBlock(
                        block_id=b_idx + 1,
                        text=para,
                        bbox=None,
                        block_type="text",
                    )
                )

            pages.append(
                PageContent(
                    page_number=idx + 1,
                    text=page_text,
                    blocks=blocks,
                    is_ocr=False,
                )
            )

        if not pages:
            pages = [PageContent(page_number=1, text=content.strip(), blocks=[], is_ocr=False)]

        return Document(
            document_id=document_id,
            source_path=str(path.resolve()),
            pages=pages,
            file_type="txt",
            metadata={"file_size_bytes": path.stat().st_size},
        )

    def _load_docx(self, path: Path, document_id: str) -> Document:
        """Load DOCX document using python-docx."""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required to load DOCX files. "
                "Install it via `pip install python-docx`."
            )

        doc = docx.Document(path)
        blocks: List[TextBlock] = []
        full_paras: List[str] = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                full_paras.append(text)
                blocks.append(
                    TextBlock(
                        block_id=idx + 1,
                        text=text,
                        bbox=None,
                        block_type="text",
                    )
                )

        full_text = "\n\n".join(full_paras)
        page = PageContent(
            page_number=1,
            text=full_text,
            blocks=blocks,
            is_ocr=False,
            metadata={"total_paragraphs": len(full_paras)},
        )

        return Document(
            document_id=document_id,
            source_path=str(path.resolve()),
            pages=[page],
            file_type="docx",
            metadata={"file_size_bytes": path.stat().st_size},
        )
